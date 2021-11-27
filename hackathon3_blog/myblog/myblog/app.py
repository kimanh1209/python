from flask import Flask, render_template, request, redirect, flash, g, make_response
from docx import Document
from pymysql import connect, cursors
import datetime

app = Flask(__name__)
app.secret_key = "kimanhphamsecretkey"

db_config = {
    "host": "localhost",
    "user": "root",
    "password": "Xinchao@@12",
    "database": "blog_mysql",
    "cursorclass": cursors.DictCursor
}


def connect_db():
    conn = connect(**db_config)
    cur = conn.cursor()
    g.conn = conn
    g.cur = cur


def close_db():
    conn = g.get("conn")
    cur = g.get("cur")

    if conn and cur:
        cur.close()
        conn.close()


@app.before_request
def before_request():
    connect_db()
    user_id = request.cookies.get('user_id')

    if user_id:
        g.cur.execute("SELECT * FROM users WHERE id = %s", user_id)
        g.user = g.cur.fetchone()
        # print(g.user.username)
    # print(g)


@app.teardown_request
def teardown_request(Exeption):
    close_db()


# @app.route("/")
# def main():
#     print('aaa')
#     return render_template("index.html")

@app.route('/about')
def showAbout():
    return render_template('about.html')

@app.route('/login')
def showLogin():
    return render_template('login.html')

@app.route('/help')
def showHelp():
    return render_template('help.html')

@app.route('/register')
def showRegister():
    return render_template('register.html')


@app.route('/register', methods=["POST"])
def register():
    username = request.form.get("username")
    password = request.form.get("newPass")
    repassword = request.form.get("reNewPass")
    if not username or not password or not repassword:
        flash("Username, password and repassword cannot be empty")
    elif len(username.strip()) < 4:
        flash("Username must have at least 4 characters")
    elif len(password) < 6:
        flash("Password must have at least 6 characters")
    elif password != repassword:
        flash("Re-Password is not correct")
    else:
        sql = '''
            SELECT * FROM users
            WHERE username = %s
        '''
        g.cur.execute(sql, username)
        user = g.cur.fetchone()
        if user:
            flash("Account already exists")
            return redirect("/login")

        sql = '''
            INSERT INTO users (username, password)
            VALUES (%s, %s)
        '''
        g.cur.execute(sql, (username.strip(), password))
        g.conn.commit()
        if g.cur.lastrowid:
            flash("Successful registration, please signin")
            return redirect("/login")
        else:
            flash("Unable to register at the moment, please try again later")
    # return render_template('register.html')
    return redirect("/register")


@app.post("/login")
def login():
    username = request.form.get("username")
    password = request.form.get("password")

    if not username or not password:
        flash("Username and password cannot be empty")
        return redirect("/login")
    else:
        sql = '''
            SELECT * FROM users
            WHERE username = %s
        '''
        g.cur.execute(sql, username)
        user = g.cur.fetchone()

        if not user or user["password"] != password:
            flash("Username or password are incorrect")
            return redirect("/login")
        else:
            response = make_response(redirect("/"))
            response.set_cookie("user_id", str(user["id"]))
            return response


@app.get("/logout")
def logout():
    response = make_response(redirect("/"))
    response.set_cookie("user_id", "")
    return response


@app.route("/", methods=["GET"])
def homepage():
    sql = '''
        SELECT * FROM posts
        JOIN users ON posts.author = users.id
        ORDER BY created DESC
    '''
    g.cur.execute(sql)

    posts = g.cur.fetchall()
    # print(posts)
    return render_template("index.html", posts=posts)


@app.route("/newpost", methods=["POST"])
def new_post():
    if not g.user:
        flash("You must be signin")
        return redirect("/")

    title = ""
    content = request.form.get("content")
    postTime = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    if content:
        sql = '''
            INSERT INTO posts (title, content, author, created)
            VALUES (%s, %s, %s, %s)
        '''
        g.cur.execute(sql, (title, content, g.user['id'], postTime))
        g.conn.commit()
        # flash("Post created!")
    else:
        flash("Content cannot be empty")

    return redirect("/")


@app.route("/post", methods=["GET"])
def render_post():
    post_id = request.args.get('post_id', default = 0, type = int)
    sql = '''
        SELECT * FROM posts
        WHERE id = %s
    '''
    g.cur.execute(sql, post_id)
    post = g.cur.fetchone()
    return render_template('post.html', post=post)


@app.route("/post", methods=["POST"])
def edit_post():
    post_id = request.args.get('post_id', default = 0, type = int)
    title = ""
    content = request.form.get("content")
    print(content)
    postTime = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    if content:
        sql = '''
            SELECT * FROM posts
            WHERE id = %s
        '''
        g.cur.execute(sql, post_id)
        post = g.cur.fetchone()

        if not post:
            flash("Post does not exists")
        elif post['author'] != g.user['id']:
            flash("You cannot edit this post")
        else:
            sql = '''
                UPDATE posts
                SET title = %s,
                    content = %s,
                    created = %s
                WHERE
                    id = %s
            '''

            g.cur.execute(sql, (title, content, postTime, post_id))
            g.conn.commit()

        return redirect("/")
    else:
        flash("Content cannot be empty")
        return redirect(request.url)

@app.route("/delete", methods=["POST"])
def delete_post():
    post_id = request.args.get('post_id', default = 0, type = int)
    sql = '''
        SELECT * FROM posts
        WHERE id = %s
    '''
    g.cur.execute(sql, post_id)
    post = g.cur.fetchone()

    if not post:
        flash("Post does not exists")
    elif post['author'] != g.user['id']:
        flash("You cannot edit this post")
    else:
        sql = '''
            DELETE FROM posts
            WHERE id = %s
        '''

        g.cur.execute(sql, (post_id))
        g.conn.commit()
    return redirect("/")

@app.route("/help", methods=["POST"])
def newTicket():
    title = request.form.get("title")
    content = request.form.get("content")
    username = request.form.get("username")
    postTime = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
    
    if title and content:
        fileName  = f"{postTime}_{username}_ticket.docx"
        filePath = f"./myblog/static/files/{fileName}"
        doc = Document()
        doc.add_heading(f"Ticket of {username}", level=0)
        doc.add_paragraph(title)
        doc.add_paragraph(content)
        doc.save(filePath)
        # print(filePath)
        return render_template("./help.html", file_created=True,filename=fileName)
        
    else:
        flash("Title and content cannot be empty")
        return render_template("./help.html", file_created=False,filename=None)


if __name__ == "__main__":
    app.run(port=5000)
